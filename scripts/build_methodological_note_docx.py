#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "docs" / "Methodological_Note_Inductive_Term_Audit.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    run._r.append(instr)

    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Page ")
    add_field(footer, "PAGE")


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Methodological Note")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(
        "Deductive and inductive term calibration for Erowid bad trip report analysis"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [2200, 7040])
    rows = [
        ("Project", "Erowid bad trip reports: phenomenology, serious adverse events, and forensic/legal screening"),
        ("Purpose", "Shared methodological protocol for collaborator review and supervised refinement"),
        ("Corpus level", "Report-level analysis after deduplication by Erowid ExpID"),
        ("Status", "Working methodological draft; suitable for collaborative review before manuscript finalization"),
    ]
    for row, (key, value) in zip(meta.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_FILL)
        row.cells[0].paragraphs[0].add_run(key).bold = True
        row.cells[1].paragraphs[0].add_run(value)


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_two_col_table(doc: Document, headers: tuple[str, str], rows: list[tuple[str, str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_FILL)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
    for left, right in rows:
        row = table.add_row()
        row.cells[0].paragraphs[0].add_run(left).bold = True
        row.cells[1].paragraphs[0].add_run(right)


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("1. Objective", level=1)
    add_para(
        doc,
        "This methodological note specifies how we distinguish deductive lexical screening from corpus-derived inductive term discovery, and how both layers are used to calibrate the analysis of Erowid bad trip reports. The protocol is designed to support collaborative review, improve transparency, reduce false positives, and prepare the dataset for later supervised multilabel modeling.",
    )
    add_para(
        doc,
        "The key principle is that corpus-derived terms are not treated as final labels. They are diagnostic signals used to audit the cohorts produced by dictionary-based screening. Unexpected terms must therefore trigger review, not immediate interpretation.",
    )

    doc.add_heading("2. Data Structure and Unit of Analysis", level=1)
    add_para(
        doc,
        "The raw scrape contains source rows rather than unique reports. Because a single Erowid report can appear across multiple substance or experience-category pages, all report-level analyses are performed after deduplication by Erowid ExpID. The local current corpus contains 34,848 source rows and 8,011 unique reports after deduplication.",
    )
    add_bullets(
        doc,
        [
            "Source rows are used only to reconstruct the local corpus and preserve category provenance.",
            "All screening, modeling, and validation steps are performed at the unique-report level unless explicitly stated otherwise.",
            "Full texts and snippets remain local and are not committed to the public repository.",
            "Aggregated term inventories and methodological code can be shared because they do not redistribute complete reports.",
        ],
    )

    doc.add_heading("3. Deductive Dictionary Layer", level=1)
    add_para(
        doc,
        "The deductive layer consists of pre-specified dictionaries derived from the research question, existing phenomenological instruments and constructs, adverse-event concepts, forensic/legal categories, and iterative identification of known false positives. This layer is intentionally high sensitivity: its purpose is to capture candidate passages and cohorts rather than to produce final validated prevalence estimates.",
    )
    add_two_col_table(
        doc,
        ("Dictionary family", "Role in the analysis"),
        [
            ("Phenomenology", "Screens for candidate reports involving fear of death, ego dissolution, mystical experience, paranoia, entity encounters, trauma re-experiencing, autobiographical insight, religious feeling, interoceptive threat, loss of control, surrender, and integration."),
            ("Serious adverse events", "Screens for candidate medical, behavioral, self-harm, injury, rescue, and life-threatening events."),
            ("Forensic/legal medicine", "Screens for candidate police contact, arrest, custody, court, charges, violence, impaired driving, endangerment, involuntary psychiatric hold, and related legal outcomes."),
            ("Contextual flags", "Detects signs of negation, hypothesis, fear, analogy, past history, and marker-specific false positives."),
        ],
        [2700, 6660],
    )
    add_para(
        doc,
        "The full transparent dictionary inventory is exported with `python3 scripts/export_keyword_inventory.py` to `outputs/tables/keyword_inventory_all.csv`.",
    )

    doc.add_heading("4. Inductive Corpus-Derived Term Layer", level=1)
    add_para(
        doc,
        "The inductive layer starts from the language of the reports themselves. It extracts unigrams, bigrams, and trigrams from the deduplicated corpus, then estimates which terms are frequent overall and which terms are specifically overrepresented in a given cohort compared with the remainder of the corpus.",
    )
    add_para(
        doc,
        "This step is implemented with `python3 scripts/extract_corpus_derived_terms.py`, which writes `outputs/tables/corpus_derived_terms.csv` and `outputs/tables/corpus_derived_terms_summary.csv`.",
    )
    add_two_col_table(
        doc,
        ("Column", "Interpretation"),
        [
            ("comparison_type", "The comparison family, e.g. overall, target_group, phenomenology_marker, serious_event_marker, forensic_legal_marker."),
            ("cohort", "The screened group whose vocabulary is being compared against the background corpus."),
            ("term", "A corpus-derived unigram, bigram, or trigram."),
            ("cohort_doc_pct", "Percentage of reports in the cohort containing the term."),
            ("background_doc_pct", "Percentage of reports outside the cohort containing the term."),
            ("specificity_log_odds", "Specificity score comparing term presence inside vs. outside the cohort."),
        ],
        [2500, 6860],
    )

    doc.add_heading("5. Calibration Logic", level=1)
    add_para(
        doc,
        "The inductive layer is used as a quality-control layer over the deductive dictionaries. A term that is overrepresented in a screened cohort does not automatically belong to that construct. It indicates that the term is associated with the cohort as currently defined. The association can reflect a true theme, a contextual co-occurrence, a subgroup, a preprocessing artifact, or a false-positive dictionary rule.",
    )
    add_numbered(
        doc,
        [
            "Apply the pre-specified deductive dictionary to the deduplicated corpus.",
            "Construct screen-positive cohorts for each marker or composite category.",
            "Extract corpus-derived n-grams and compute specificity scores for each cohort.",
            "Review high-specificity terms for conceptual coherence and possible artifacts.",
            "For unexpected or incoherent terms, inspect a small number of local source reports or snippets.",
            "If the term reveals a false-positive rule, tighten the corresponding pattern or add a marker-specific exclusion.",
            "Regenerate the screening tables, human-review files, and corpus-derived term tables.",
            "Repeat until the major cohorts are stable enough for human annotation and quantitative description.",
        ],
    )

    doc.add_heading("6. Decision Rules for Collaborator Review", level=1)
    add_para(
        doc,
        "During collaborative review, terms in the inductive file should be triaged into one of the following categories. This classification is a methodological audit step, not a final scientific finding.",
    )
    add_two_col_table(
        doc,
        ("Review category", "Operational meaning"),
        [
            ("Coherent term", "The term is conceptually consistent with the cohort and may help name or interpret it."),
            ("Unanticipated but plausible term", "The term suggests a theme not fully captured by the current codebook and should be considered for annotation."),
            ("Ambiguous term", "The term requires local reading of reports because it can carry multiple meanings."),
            ("Likely false-positive signal", "The term indicates that the dictionary rule may be capturing irrelevant contexts."),
            ("Boilerplate or scraping artifact", "The term likely comes from page furniture, advertisements, metadata, or repeated non-narrative material."),
            ("Low interpretive value", "The term is frequent or specific but too generic to support interpretation."),
        ],
        [2600, 6760],
    )

    doc.add_heading("7. Human Annotation and Model Development", level=1)
    add_para(
        doc,
        "The calibrated dictionaries and inductive audits are preparatory steps for human annotation and supervised modeling. They should not be treated as replacements for human coding. The recommended modeling sequence is a hybrid workflow: transparent screening, inductive audit, snippet-level and report-level annotation, baseline supervised multilabel classification, and error analysis.",
    )
    add_bullets(
        doc,
        [
            "Annotation should distinguish confirmed events from fears, hypotheses, analogies, past history, and metaphors.",
            "Phenomenological labels should be multilabel and may include valence and narrative phase.",
            "A first supervised baseline should use TF-IDF n-grams with logistic regression or a linear SVM, evaluated label by label.",
            "Embedding-based clustering and topic modeling should be used exploratorily to find missed themes and refine the codebook.",
            "Final estimates should report dictionary-screened counts separately from human-validated or model-estimated labels.",
        ],
    )

    doc.add_heading("8. Transparency and Reproducibility", level=1)
    add_para(
        doc,
        "For publication, the methods section should clearly separate the four layers: raw source rows, deduplicated unique reports, deductive dictionary screening, and inductive term auditing. The dictionary inventory and corpus-derived term inventory should be treated as methodological supplements.",
    )
    add_two_col_table(
        doc,
        ("Artifact", "Purpose"),
        [
            ("keyword_inventory_all.csv", "Documents all deductive dictionaries, rule families, contextual flags, and known false-positive patterns."),
            ("corpus_derived_terms.csv", "Documents terms emerging from the corpus overall and within screened cohorts."),
            ("human_review_workbook.xlsx", "Local review file for human validation; contains snippets and must not be publicly redistributed."),
            ("modeling_and_transparency_protocol.md", "Repository-level methodological protocol for the complete NLP workflow."),
            ("deductive_vs_inductive_terms.md", "Repository-level explanation of dictionary screening vs. corpus-derived term auditing."),
        ],
        [3100, 6260],
    )

    doc.add_heading("9. Suggested Manuscript Language", level=1)
    add_para(
        doc,
        "We used a two-stage lexical strategy combining deductive screening and inductive term auditing. First, high-sensitivity dictionaries were specified for phenomenological, serious adverse-event, and forensic/legal constructs. These dictionaries were used to construct candidate cohorts, not final validated outcomes. Second, corpus-derived unigrams, bigrams, and trigrams were extracted from the deduplicated report corpus and scored for specificity within each screen-positive cohort relative to the remaining corpus. Terms overrepresented in each cohort were inspected as a diagnostic layer to identify coherent themes, unanticipated concepts, ambiguous language, boilerplate artifacts, and false-positive dictionary rules. When clear false-positive patterns were identified, the corresponding dictionary rule was tightened and the pipeline was rerun. Human validation and later supervised multilabel classification are planned as subsequent steps to estimate validated construct prevalence.",
    )

    doc.add_heading("10. Open Tasks for the Research Team", level=1)
    add_bullets(
        doc,
        [
            "Review the highest-specificity terms for priority cohorts and classify them using the audit categories above.",
            "Identify terms that suggest missing phenomenological labels or sublabels.",
            "Identify terms that reveal false-positive dictionary rules or boilerplate contamination.",
            "Select a stratified annotation sample for serious adverse events, forensic/legal events, and core phenomenological labels.",
            "Define the minimum annotation set required before training a supervised multilabel model.",
        ],
    )

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "Note: This document intentionally describes the calibration procedure without relying on case-specific examples from the corpus, so it can be shared as a general methodological protocol."
    )
    run.italic = True
    run.font.color.rgb = MUTED

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_doc()
